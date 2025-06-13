"""
Flask Inventory Slip Generator - Web application for generating inventory slips
from CSV and JSON data with support for Bamboo and Cultivera formats.
"""

# Standard library imports
import os
import sys
import json
import datetime
import socket
import ssl
import base64
import hmac
import hashlib
import logging
import threading
import tempfile
import urllib.request
import urllib.error
import uuid
import re
import webbrowser
import time
from functools import wraps
from io import BytesIO
import zlib

# Third-party imports
from flask import (
    Flask, 
    render_template, 
    request, 
    redirect, 
    url_for, 
    flash, 
    jsonify, 
    session, 
    send_file, 
    send_from_directory
)
import requests
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docxcompose.composer import Composer
import configparser
from werkzeug.utils import secure_filename

# Local imports
from src.utils.document_handler import DocumentHandler
from src.ui.app import InventorySlipGenerator


import sys
def compress_session_data(data):
    """Compress data for session storage"""
    if isinstance(data, str):
        compressed = zlib.compress(data.encode('utf-8'))
    else:
        compressed = zlib.compress(json.dumps(data).encode('utf-8'))
    return base64.b64encode(compressed).decode('utf-8')

def decompress_session_data(compressed_data):
    """Decompress data from session storage"""
    if not compressed_data:
        return None
    try:
        decompressed = zlib.decompress(base64.b64decode(compressed_data))
        return json.loads(decompressed)
    except Exception as e:
        logger.error(f"Failed to decompress session data: {str(e)}")
        return None

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Constants
CONFIG_FILE = os.path.expanduser("~/inventory_generator_config.ini")

def get_downloads_dir():
    """Get the default Downloads directory for both Windows and Mac"""
    try:
        if sys.platform == "win32":
            # First try Windows known folder path
            import winreg
            from ctypes import windll, wintypes
            CSIDL_PERSONAL = 5  # Documents
            SHGFP_TYPE_CURRENT = 0  # Get current path
            buf = wintypes.create_unicode_buffer(wintypes.MAX_PATH)
            windll.shell32.SHGetFolderPathW(None, CSIDL_PERSONAL, None, SHGFP_TYPE_CURRENT, buf)
            documents = buf.value
            
            # Try registry next
            try:
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER, 
                    r"SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders") as key:
                    downloads = winreg.QueryValueEx(key, "{374DE290-123F-4565-9164-39C4925E467B}")[0]
                return downloads
            except:
                # Fall back to Documents\Downloads
                return os.path.join(documents, "Downloads")
        else:
            # macOS and Linux
            return os.path.join(os.path.expanduser("~"), "Downloads")
    except:
        # Ultimate fallback - user's home directory
        return os.path.expanduser("~")

# Update the constants
DEFAULT_SAVE_DIR = get_downloads_dir()
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), "inventory_generator", "uploads")

# Ensure directories exist with proper permissions
try:
    os.makedirs(DEFAULT_SAVE_DIR, exist_ok=True)
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
except Exception as e:
    logger.error(f"Error creating directories: {str(e)}")
    # Fall back to temp directory if needed
    if not os.path.exists(DEFAULT_SAVE_DIR):
        DEFAULT_SAVE_DIR = tempfile.gettempdir()

APP_VERSION = "2.0.0"
ALLOWED_EXTENSIONS = {'csv', 'json', 'docx'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16 MB max upload size

# Add new constants for API configuration
API_CONFIGS = {
    'bamboo': {
        'base_url': 'https://api-trace.getbamboo.com/shared/manifests',
        'version': 'v1',
        'auth_type': 'bearer'
    },
    'cultivera': {
        'base_url': 'https://api.cultivera.com/api',
        'version': 'v1',
        'auth_type': 'basic'
    },
    'growflow': {
        'base_url': 'https://api.growflow.com',
        'version': 'v2',
        'auth_type': 'oauth2'
    }
}

# Initialize Flask application
app = Flask(__name__,
    static_url_path='',
    static_folder='static',
    template_folder='templates'
)
# Use a fixed secret key for development to preserve session data
app.secret_key = 'your-fixed-development-secret-key'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Helper function to get resource path (for templates)
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

# Load configurations or create default
def load_config():
    config = configparser.ConfigParser()
    
    # Default configurations
    config['PATHS'] = {
        'template_path': os.path.join(os.path.dirname(__file__), "templates/documents/InventorySlips.docx"),
        'output_dir': DEFAULT_SAVE_DIR,  # Use the new DEFAULT_SAVE_DIR
        'recent_files': '',
        'recent_urls': ''
    }
    
    config['SETTINGS'] = {
        'items_per_page': '4',
        'auto_open': 'true',
        'theme': 'dark',
        'font_size': '12'
    }
    
    # Load existing config if it exists
    if os.path.exists(CONFIG_FILE):
        config.read(CONFIG_FILE)
    else:
        # Create config file with defaults
        with open(CONFIG_FILE, 'w') as f:
            config.write(f)
    
    return config

def save_config(config):
    with open(CONFIG_FILE, 'w') as f:
        config.write(f)

# Helper to adjust font sizes after rendering
def adjust_table_font_sizes(doc_path):
    """
    Post-process a DOCX file to dynamically adjust font size inside table cells based on thresholds.
    """
    thresholds = [
        (30, 12),   # <=30 chars → 12pt
        (45, 10),   # <=45 chars → 10pt
        (60, 8),    # <=60 chars → 8pt
        (float('inf'), 7)  # >60 chars → 7pt
    ]

    def get_font_size(text_len):
        for limit, size in thresholds:
            if text_len <= limit:
                return size
        return 7  # Fallback

    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue

                    # If line is Product Name (first line), force 10pt
                    if paragraph == cell.paragraphs[0]:
                        font_size = 10
                    else:
                        font_size = get_font_size(len(text))

                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)

    doc.save(doc_path)

# Open files after saving
def open_file(path):
    """Open files using the default system application"""
    try:
        if sys.platform == "win32":
            os.startfile(path)
        elif sys.platform == "darwin":  # macOS
            os.system(f'open "{path}"')
        else:  # linux variants
            os.system(f'xdg-open "{path}"')
    except Exception as e:
        logger.error(f"Error opening file: {e}")
        flash(f"Error opening file: {e}", "error")

# Split records into chunks
def chunk_records(records, chunk_size=4):
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

# Check if file extension is allowed
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Process and save inventory slips
def run_full_process_inventory_slips(selected_df, config, status_callback=None, progress_callback=None):
    if selected_df.empty:
        if status_callback:
            status_callback("Error: No data selected.")
        return False, "No data selected."

    try:
        # Get settings from config
        items_per_page = int(config['SETTINGS'].get('items_per_page', '4'))
        template_path = config['PATHS'].get('template_path')
        if not template_path or not os.path.exists(template_path):
            template_path = os.path.join(os.path.dirname(__file__), "templates/documents/InventorySlips.docx")
            if not os.path.exists(template_path):
                raise ValueError(f"Template file not found at: {template_path}")
        
        if status_callback:
            status_callback("Processing data...")

        records = selected_df.to_dict(orient="records")
        pages = []

        # Process records in chunks of 4 (or configured size)
        total_chunks = (len(records) + items_per_page - 1) // items_per_page
        current_chunk = 0

        for chunk in chunk_records(records, items_per_page):
            current_chunk += 1
            if progress_callback:
                progress = (current_chunk / total_chunks) * 50
                progress_callback(int(progress))

            if status_callback:
                status_callback(f"Generating page {current_chunk} of {total_chunks}...")

            try:
                tpl = DocxTemplate(template_path)
                context = {}

                # Fill context with records - modified vendor handling
                for idx, record in enumerate(chunk, 1):
                    # Get vendor info, using full vendor name if available
                    vendor_name = record.get("Vendor", "")
                    # If vendor is in format "license - name", extract just the name
                    if " - " in vendor_name:
                        vendor_name = vendor_name.split(" - ")[1]
                    
                    context[f"Label{idx}"] = {
                        "ProductName": record.get("Product Name*", ""),
                        "Barcode": record.get("Barcode*", ""),
                        "AcceptedDate": record.get("Accepted Date", ""),
                        "QuantityReceived": record.get("Quantity Received*", ""),
                        "Vendor": vendor_name or "Unknown Vendor",  # Only use Unknown if empty
                        "ProductType": record.get("Product Type*", "")
                    }

                # Fill remaining slots with empty values
                for i in range(len(chunk) + 1, items_per_page + 1):
                    context[f"Label{i}"] = {
                        "ProductName": "",
                        "Barcode": "",
                        "AcceptedDate": "",
                        "QuantityReceived": "",
                        "Vendor": "",
                        "ProductType": ""
                    }

                # Render template with context
                tpl.render(context)
                
                # Save to BytesIO
                output = BytesIO()
                tpl.save(output)
                pages.append(Document(output))

            except Exception as e:
                raise ValueError(f"Error generating page {current_chunk}: {e}")

        if not pages:
            return False, "No documents generated."

        # Combine pages
        if status_callback:
            status_callback("Combining pages...")

        master = pages[0]
        composer = Composer(master)
        for i, doc in enumerate(pages[1:]):
            if progress_callback:
                progress = 50 + ((i + 1) / len(pages[1:])) * 40
                progress_callback(int(progress))
            composer.append(doc)

        # Save final document
        now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        outname = f"inventory_slips_{now}.docx"
        outpath = os.path.join(config['PATHS']['output_dir'], outname)

        if status_callback:
            status_callback("Saving document...")

        master.save(outpath)

        # Adjust font sizes
        if status_callback:
            status_callback("Adjusting formatting...")
        adjust_table_font_sizes(outpath)

        if progress_callback:
            progress_callback(100)

        return True, outpath

    except Exception as e:
        if status_callback:
            status_callback(f"Error: {str(e)}")
        return False, str(e)

# Parse Bamboo transfer schema JSON
def parse_bamboo_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Get vendor information
        from_license_number = json_data.get("from_license_number", "")
        from_license_name = json_data.get("from_license_name", "")
        vendor_meta = f"{from_license_number} - {from_license_name}"
        
        # Get transfer date
        raw_date = json_data.get("est_arrival_at", "") or json_data.get("transferred_at", "")
        accepted_date = raw_date.split("T")[0] if "T" in raw_date else raw_date
        
        # Process inventory items
        items = json_data.get("inventory_transfer_items", [])
        logger.info(f"Bamboo data: found {len(items)} inventory_transfer_items")
        records = []
        
        for item in items:
            # Extract THC and CBD content from lab_result_data if available
            thc_content = ""
            cbd_content = ""
            
            lab_data = item.get("lab_result_data", {})
            if lab_data and "potency" in lab_data:
                for potency_item in lab_data["potency"]:
                    if potency_item.get("type") == "total-thc":
                        thc_content = f"{potency_item.get('value', '')}%"
                    elif potency_item.get("type") == "total-cbd":
                        cbd_content = f"{potency_item.get('value', '')}%"
            
            records.append({
                "Product Name*": item.get("product_name", ""),
                "Product Type*": item.get("inventory_type", ""),
                "Quantity Received*": item.get("qty", ""),
                "Barcode*": item.get("inventory_id", "") or item.get("external_id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": item.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Bamboo"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Bamboo transfer data: {e}")

# Parse Cultivera JSON
def parse_cultivera_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Check if Cultivera format
        if not json_data.get("data") or not isinstance(json_data.get("data"), dict):
            raise ValueError("Not a valid Cultivera format")
        
        data = json_data.get("data", {})
        manifest = data.get("manifest", {})
        
        # Get vendor information
        from_license = manifest.get("from_license", {})
        vendor_name = from_license.get("name", "")
        vendor_license = from_license.get("license_number", "")
        vendor_meta = f"{vendor_license} - {vendor_name}" if vendor_license and vendor_name else "Unknown Vendor"
        
        # Get transfer date
        created_at = manifest.get("created_at", "")
        accepted_date = created_at.split("T")[0] if "T" in created_at else created_at
        
        # Process inventory items
        items = manifest.get("items", [])
        records = []
        
        for item in items:
            # Extract product info
            product = item.get("product", {})
            
            # Extract THC and CBD content
            thc_content = ""
            cbd_content = ""
            
            test_results = item.get("test_results", [])
            if test_results:
                for result in test_results:
                    if "thc" in result.get("type", "").lower():
                        thc_content = f"{result.get('percentage', '')}%"
                    elif "cbd" in result.get("type", "").lower():
                        cbd_content = f"{result.get('percentage', '')}%"
            
            records.append({
                "Product Name*": product.get("name", ""),
                "Product Type*": product.get("category", ""),
                "Quantity Received*": item.get("quantity", ""),
                "Barcode*": item.get("barcode", "") or item.get("id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": product.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Cultivera"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Cultivera data: {e}")

def parse_growflow_data(json_data):
    """Parse GrowFlow JSON format into common fields"""
    try:
        if not ('inventory_transfer_items' in json_data and 
                'from_license_number' in json_data and 
                'from_license_name' in json_data):
            return pd.DataFrame()
        
        vendor_meta = f"{json_data.get('from_license_number', '')} - {json_data.get('from_license_name', 'Unknown Vendor')}"
        raw_date = json_data.get("est_arrival_at", "") or json_data.get("transferred_at", "")
        accepted_date = raw_date.split("T")[0] if "T" in raw_date else raw_date
        
        items = json_data.get("inventory_transfer_items", [])
        mapped_data = []
        
        for item in items:
            potency_data = item.get("lab_result_data", {}).get("potency", [])
            thc_value = next((p.get('value') for p in potency_data if p.get('type') in ["total-thc", "thc"]), 0)
            cbd_value = next((p.get('value') for p in potency_data if p.get('type') in ["total-cbd", "cbd"]), 0)
            
            mapped_item = {
                "Product Name*": item.get("product_name", ""),
                "Product Type*": item.get("inventory_type", ""),
                "Quantity Received*": item.get("qty", ""),
                "Barcode*": item.get("product_sku", "") or item.get("inventory_id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": item.get("strain_name", ""),
                "THC Content": f"{thc_value}%",
                "CBD Content": f"{cbd_value}%",
                "Source System": "GrowFlow"
            }
            mapped_data.append(mapped_item)
        
        return pd.DataFrame(mapped_data)
    
    except Exception as e:
        logger.error(f"Error parsing GrowFlow data: {str(e)}")
        return pd.DataFrame()

def parse_inventory_json(json_data):
    """
    Detects and parses JSON format accordingly
    Returns tuple of (DataFrame, format_type)
    """
    if not json_data:
        return None, "No data provided"
    
    try:
        # Parse string to JSON if needed
        if isinstance(json_data, str):
            json_data = json.loads(json_data)
            
        # Try parsing as Bamboo
        if "inventory_transfer_items" in json_data:
            return parse_bamboo_data(json_data), "Bamboo"
            
        # Try parsing as Cultivera 
        elif "data" in json_data and isinstance(json_data["data"], dict) and "manifest" in json_data["data"]:
            return parse_cultivera_data(json_data), "Cultivera"
            
        # Try parsing as GrowFlow
        elif "document_schema_version" in json_data:
            return parse_growflow_data(json_data), "GrowFlow"
            
        else:
            return None, "Unknown JSON format"
            
    except json.JSONDecodeError:
        return None, "Invalid JSON data"
    except Exception as e:
        return None, f"Error parsing data: {str(e)}"

# Process CSV data
def process_csv_data(df):
    try:
        # Strip whitespace from column names
        df.columns = [col.strip() for col in df.columns]
        logger.info(f"Original columns: {df.columns.tolist()}")
        
        # First, ensure column names are unique by adding a suffix if needed
        df.columns = [f"{col}_{i}" if df.columns.tolist().count(col) > 1 else col 
                     for i, col in enumerate(df.columns)]
        logger.info(f"Columns after ensuring uniqueness: {df.columns.tolist()}")
        
        # Map column names to expected format
        col_map = {
            "Product Name*": "Product Name*",
            "Product Name": "Product Name*",
            "Quantity Received": "Quantity Received*",
            "Quantity*": "Quantity Received*",
            "Quantity": "Quantity Received*",
            "Lot Number*": "Barcode*",
            "Barcode": "Barcode*",
            "Lot Number": "Barcode*",
            "Accepted Date": "Accepted Date",
            "Vendor": "Vendor",
            "Strain Name": "Strain Name",
            "Product Type*": "Product Type*",
            "Product Type": "Product Type*",
            "Inventory Type": "Product Type*"
        }
        
        # Now rename columns according to our mapping
        new_columns = {}
        target_counts = {}  # Keep track of how many times we've used each target name
        
        for col in df.columns:
            base_col = col.split('_')[0]  # Remove any suffix
            if base_col in col_map:
                target_name = col_map[base_col]
                # If we've seen this target name before, add a suffix
                if target_name in target_counts:
                    target_counts[target_name] += 1
                    new_columns[col] = f"{target_name}_{target_counts[target_name]}"
                else:
                    target_counts[target_name] = 0
                    new_columns[col] = target_name
            else:
                new_columns[col] = col
        
        logger.info(f"Column mapping: {new_columns}")
        df = df.rename(columns=new_columns)
        logger.info(f"Columns after renaming: {df.columns.tolist()}")
        
        # Ensure required columns exist
        required_cols = ["Product Name*", "Barcode*"]
        missing_cols = [col for col in required_cols if not any(col in c for c in df.columns)]
        
        if missing_cols:
            return None, f"CSV is missing required columns: {', '.join(missing_cols)}"
        
        # Set default values for missing columns
        if not any("Vendor" in c for c in df.columns):
            df["Vendor"] = "Unknown Vendor"
        else:
            vendor_col = next(c for c in df.columns if "Vendor" in c)
            df[vendor_col] = df[vendor_col].fillna("Unknown Vendor")
        
        if not any("Accepted Date" in c for c in df.columns):
            today = datetime.datetime.today().strftime("%Y-%m-%d")
            df["Accepted Date"] = today
        
        if not any("Product Type*" in c for c in df.columns):
            df["Product Type*"] = "Unknown"
        
        if not any("Strain Name" in c for c in df.columns):
            df["Strain Name"] = ""
        
        # Sort if possible
        try:
            sort_cols = []
            if any("Product Type*" in c for c in df.columns):
                sort_cols.append(next(c for c in df.columns if "Product Type*" in c))
            if any("Product Name*" in c for c in df.columns):
                sort_cols.append(next(c for c in df.columns if "Product Name*" in c))
            
            if sort_cols:
                df = df.sort_values(sort_cols, ascending=[True, True])
        except:
            pass  # If sorting fails, continue without sorting
        
        # Final check for duplicate columns
        if len(df.columns) != len(set(df.columns)):
            duplicates = [col for col in df.columns if df.columns.tolist().count(col) > 1]
            logger.error(f"Duplicate columns found: {duplicates}")
            return None, f"Duplicate columns found: {', '.join(duplicates)}"
        
        return df, "Success"
    
    except Exception as e:
        logger.error(f"Error in process_csv_data: {str(e)}", exc_info=True)
        return None, f"Failed to process CSV data: {e}"

# First, update the chunk size and compression level
def chunk_session_data(data, chunk_size=3000):
    """Split large data into smaller chunks with higher compression"""
    if not isinstance(data, str):
        data = json.dumps(data)
    
    # Use higher compression level (9 is highest)
    compressed = zlib.compress(data.encode('utf-8'), level=9)
    encoded = base64.b64encode(compressed).decode('utf-8')
    
    # Calculate optimal chunk size
    total_size = len(encoded)
    num_chunks = (total_size + chunk_size - 1) // chunk_size
    if num_chunks * 40 + total_size > 4000:  # Account for chunk metadata
        chunk_size = max(1000, (4000 - num_chunks * 40) // num_chunks)
    
    chunks = [encoded[i:i+chunk_size] for i in range(0, len(encoded), chunk_size)]
    return chunks

@app.route('/paste-json', methods=['POST'])
def paste_json():
    try:
        data = request.get_json()
        pasted_json = data.get('json_data', '')
        if not pasted_json:
            return jsonify({'success': False, 'message': 'No JSON data provided.'}), 400

        try:
            parsed = json.loads(pasted_json)
        except Exception as e:
            return jsonify({'success': False, 'message': f'Invalid JSON: {str(e)}'}), 400

        result_df, format_type = parse_inventory_json(parsed)
        if result_df is None or result_df.empty:
            return jsonify({'success': False, 'message': 'Could not process pasted JSON data.'}), 400

        session['df_json'] = result_df.to_json(orient='records')
        session['format_type'] = format_type
        session['raw_json'] = pasted_json

        return jsonify({'success': True, 'redirect': url_for('data_view')})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

@app.route('/upload-csv', methods=['POST'])
def upload_csv():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(url_for('index'))
    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(url_for('index'))
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        try:
            df = pd.read_csv(filepath)
            processed_df, msg = process_csv_data(df)
            if processed_df is None:
                flash(msg)
                return redirect(url_for('index'))
            session['df_json'] = processed_df.to_json(orient='records')
            session['format_type'] = 'CSV'
            session['raw_json'] = df.to_json(orient='records')
            flash('CSV uploaded and processed successfully')
            return redirect(url_for('data_view'))
        except Exception as e:
            flash(f'Failed to process CSV: {str(e)}')
            return redirect(url_for('index'))
    else:
        flash('Invalid file type')
        return redirect(url_for('index'))

# Then, update the URL loading function
@app.route('/load-url', methods=['POST'])
def load_url():
    """Handle URL loading with proper error handling"""
    try:
        url = request.form.get('url', '').strip()
        
        if not url:
            flash('Please enter a URL', 'error')
            return redirect(url_for('index'))
        
        # Clear any existing session data
        clear_chunked_data('df_json')
        clear_chunked_data('raw_json')
        session.pop('format_type', None)
        
        # Handle different URL types
        if "bamboo" in url.lower() or "getbamboo" in url.lower():
            return handle_bamboo_url(url)
        else:
            result_df, format_type = load_from_url(url)
            
            if result_df is None or result_df.empty:
                flash('Could not process data from URL', 'error')
                return redirect(url_for('index'))
            
            # Store compressed results in chunks
            logger.info("Storing URL data in chunks...")
            store_chunked_data('df_json', result_df.to_json(orient='records', default_handler=str))
            session['format_type'] = format_type
            
            flash(f'{format_type} data loaded successfully', 'success')
            return redirect(url_for('data_view'))
            
    except ValueError as ve:
        logger.error(f'URL validation error: {str(ve)}')
        flash(str(ve), 'error')
        return redirect(url_for('index'))
    except Exception as e:
        logger.error(f'Error loading URL: {str(e)}', exc_info=True)
        flash(f'Error loading data: {str(e)}', 'error')
        return redirect(url_for('index'))
    
def handle_bamboo_url(url):
    try:
        result_df, format_type = load_from_url(url)
        if result_df is None or result_df.empty:
            flash('Could not process Bamboo data from URL', 'error')
            return redirect(url_for('index'))
        store_chunked_data('df_json', result_df.to_json(orient='records', default_handler=str))
        session['format_type'] = format_type
        flash(f'{format_type} data loaded successfully', 'success')
        return redirect(url_for('data_view'))
    except Exception as e:
        logger.error(f'Error loading Bamboo URL: {str(e)}', exc_info=True)
        flash(f'Error loading Bamboo data: {str(e)}', 'error')
        return redirect(url_for('index'))

def load_from_url(url):
    """Download JSON or CSV data from a URL and return as DataFrame and format_type."""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        content_type = response.headers.get('Content-Type', '').lower()
        if 'application/json' in content_type or url.lower().endswith('.json'):
            data = response.json()
            df, format_type = parse_inventory_json(data)
            return df, format_type
        elif 'text/csv' in content_type or url.lower().endswith('.csv'):
            df = pd.read_csv(BytesIO(response.content))
            df, msg = process_csv_data(df)
            return df, 'CSV'
        else:
            # Try to parse as JSON first, then CSV
            try:
                data = response.json()
                df, format_type = parse_inventory_json(data)
                return df, format_type
            except Exception:
                try:
                    df = pd.read_csv(BytesIO(response.content))
                    df, msg = process_csv_data(df)
                    return df, 'CSV'
                except Exception as e:
                    raise ValueError(f"Unsupported data format or failed to parse: {e}")
    except Exception as e:
        raise ValueError(f"Failed to load data from URL: {e}")
    


# Update data view to handle chunked data properly
@app.route('/data-view')
def data_view():
    try:
        # Get chunked data from session
        df_json = get_chunked_data('df_json')
        format_type = session.get('format_type')

        if df_json is None:
            flash('No data available. Please load data first.')
            return redirect(url_for('index'))

        try:
            if isinstance(df_json, list):
                df = pd.DataFrame(df_json)
            else:
                df = pd.read_json(df_json, orient='records')
        except Exception as e:
            logger.error(f"Error parsing JSON data: {str(e)}")
            flash('Error loading data. Please try again.')
            return redirect(url_for('index'))
        
        # Format data for template
        products = []
        for idx, row in df.iterrows():
            product = {
                'id': idx,
                'name': str(row.get('Product Name*', '')),
                'strain': str(row.get('Strain Name', '')),
                'sku': str(row.get('Barcode*', '')),
                'quantity': str(row.get('Quantity Received*', '')),
                'source': format_type or 'Unknown'
            }
            products.append(product)

        # Load configuration
        config = load_config()

        return render_template(
            'data_view.html',
            products=products,
            format_type=format_type,
            theme=config['SETTINGS'].get('theme', 'dark'),
            version=APP_VERSION
        )
    except Exception as e:
        logger.error(f'Error in data_view: {str(e)}', exc_info=True)
        flash('Error loading data. Please try again.')
        return redirect(url_for('index'))

@app.route('/generate-slips', methods=['POST'])
def generate_slips():
    try:
        # Get selected products
        selected_indices = request.form.getlist('selected_indices[]')
        
        if not selected_indices:
            flash('No products selected.')
            return redirect(url_for('data_view'))
        
        # Convert indices to integers
        selected_indices = [int(idx) for idx in selected_indices]
        
        # Load data from session
        df_json = session.get('df_json', None)
        
        if df_json is None:
            flash('No data available. Please load data first.')
            return redirect(url_for('index'))
        
        # Convert JSON to DataFrame
        df = pd.read_json(df_json, orient='records')
        
        # Get only selected rows
        selected_df = df.iloc[selected_indices].copy()
        
        # Load configuration
        config = load_config()
        
        # Generate the file
        status_messages = []
        progress_values = []
        
        def status_callback(msg):
            status_messages.append(msg)
        
        def progress_callback(value):
            progress_values.append(value)
        
        success, result = run_full_process_inventory_slips(
            selected_df,
            config,
            status_callback,
            progress_callback
        )
        
        if success:
            # Return the file for download
            return send_file(
                result,
                as_attachment=True,
                download_name=os.path.basename(result),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            flash(f'Failed to generate inventory slips: {result}')
            return redirect(url_for('data_view'))
    
    except Exception as e:
        flash(f'Error generating slips: {str(e)}')
        return redirect(url_for('data_view'))

@app.route('/show-result')
def show_result():
    # Get output file path from session
    output_file = session.get('output_file', None)
    
    if not output_file or not os.path.exists(output_file):
        flash('No output file available.')
        return redirect(url_for('index'))
    
    # Get filename for display
    filename = os.path.basename(output_file)
    
    # Load configuration
    config = load_config()
    
    return render_template(
        'result.html',
        filename=filename,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/download-file')
def download_file():
    # Get output file path from session
    output_file = session.get('output_file', None)
    
    if not output_file or not os.path.exists(output_file):
        flash('No output file available.')
        return redirect(url_for('index'))
    
    # Return the file for download
    return send_file(output_file, as_attachment=True)

@app.route('/settings', methods=['GET', 'POST'])
def settings():
    config = load_config()
    
    if request.method == 'POST':
        # Update settings from form
        if 'items_per_page' in request.form:
            config['SETTINGS']['items_per_page'] = request.form['items_per_page']
        
        if 'theme' in request.form:
            config['SETTINGS']['theme'] = request.form['theme']
        
        if 'api_key' in request.form:
            if 'API' not in config:
                config['API'] = {}
            config['API']['bamboo_key'] = request.form['api_key']
        
        if 'output_dir' in request.form:
            output_dir = request.form['output_dir']
            if output_dir and os.path.exists(output_dir):
                config['PATHS']['output_dir'] = output_dir
        
        # Save updated config
        save_config(config)
        flash('Settings saved successfully')
        return redirect(url_for('index'))
    
    return render_template(
        'settings.html',
        config=config,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/view-json')
def view_json():
    raw_json = session.get('raw_json', None)
    format_type = session.get('format_type', None)
    
    if raw_json is None:
        flash('No JSON data available.')
        return redirect(url_for('index'))
    
    # Load configuration
    config = load_config()
    
    return render_template(
        'view_json.html',
        raw_json=raw_json,
        format_type=format_type,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/clear-data')
def clear_data():
    # Clear chunked session data
    clear_chunked_data('df_json')
    clear_chunked_data('raw_json')
    session.pop('format_type', None)
    session.pop('output_file', None)
    
    flash('Data cleared successfully')
    return redirect(url_for('index'))

@app.route('/about')
def about():
    config = load_config()
    return render_template(
        'about.html',
        version=APP_VERSION,
        theme=config['SETTINGS'].get('theme', 'dark')
    )

@app.route('/')
def index():
    config = load_config()
    return render_template(
        'index.html',
        config=config,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/search-json-or-api', methods=['POST'])
def search_json_or_api():
    user_input = request.form.get('search_input', '').strip()
    if not user_input:
        flash('Please enter JSON data or an API URL.')
        return redirect(url_for('index'))

    # Try to detect if input is a URL
    if user_input.startswith('http://') or user_input.startswith('https://'):
        try:
            with urllib.request.urlopen(user_input) as resp:
                data = json.loads(resp.read().decode())
            result_df, format_type = parse_inventory_json(data)
            if result_df is None or result_df.empty:
                flash(f'Could not process data from URL.')
                return redirect(url_for('index'))
            session['df_json'] = result_df.to_json(orient='records')
            session['format_type'] = format_type
            session['raw_json'] = json.dumps(data)
            flash(f'{format_type} data loaded successfully from URL')
            return redirect(url_for('data_view'))
        except Exception as e:
            flash(f'Failed to load data from URL: {str(e)}')
            return redirect(url_for('index'))
    else:
        # Try to parse as JSON
        try:
            data = json.loads(user_input)
            result_df, format_type = parse_inventory_json(data)
            if result_df is None or result_df.empty:
                flash(f'Could not process pasted JSON data.')
                return redirect(url_for('index'))
            session['df_json'] = result_df.to_json(orient='records')
            session['format_type'] = format_type
            session['raw_json'] = user_input
            flash(f'{format_type} data imported successfully')
            return redirect(url_for('data_view'))
        except Exception as e:
            flash(f'Failed to import JSON data: {str(e)}')
            return redirect(url_for('index'))

# Error handlers
@app.errorhandler(404)
def page_not_found(e):
    config = load_config()
    return render_template('404.html', theme=config['SETTINGS'].get('theme', 'dark')), 404

@app.errorhandler(500)
def server_error(e):
    config = load_config()
    return render_template('500.html', theme=config['SETTINGS'].get('theme', 'dark')), 500

def validate_docx(file_path):
    """Validate the generated DOCX file"""
    try:
        doc = Document(file_path)
        # Try to access content to verify document is readable
        _ = doc.paragraphs
        _ = doc.tables
        return True
    except Exception as e:
        logger.error(f"Document validation failed: {str(e)}")
        return False

@app.route('/select-directory', methods=['POST'])
def select_directory():
    selected_dir = request.form.get('directory')
    if selected_dir and os.path.exists(selected_dir):
        config = load_config()
        config['PATHS']['output_dir'] = selected_dir
        save_config(config)
        return jsonify({
            'success': True,
            'selected_dir': selected_dir,
            'message': 'Output directory updated successfully'
        })
    return jsonify({
        'success': False,
        'message': 'Invalid directory selected'
    }), 400

import subprocess
from flask import jsonify

@app.route('/open_downloads')
def open_downloads():
    downloads_dir = get_downloads_dir()  # This function returns the appropriate downloads folder
    try:
        if sys.platform == "darwin":  # macOS
            subprocess.Popen(["open", downloads_dir])
        elif sys.platform == "win32":  # Windows
            os.startfile(downloads_dir)
        else:  # Linux and other OSes
            subprocess.Popen(["xdg-open", downloads_dir])
        return jsonify(success=True)
    except Exception as e:
        return jsonify(success=False, message=str(e))

def create_api_signature(secret_key, message):
    """Create HMAC signature for API requests"""
    key = secret_key.encode('utf-8')
    message = message.encode('utf-8')
    signature = hmac.new(key, message, hashlib.sha256).hexdigest()
    return signature

def require_api_key(f):
    """Decorator to require API key for certain routes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        config = load_config()
        api_type = request.args.get('api_type', 'bamboo')
        
        if 'API' not in config or f'{api_type}_key' not in config['API']:
            flash(f'No API key configured for {api_type}. Please add it in settings.')
            return redirect(url_for('settings'))
        return f(*args, **kwargs)
    return decorated_function

class APIClient:
    def __init__(self, api_type, config):
        self.api_type = api_type
        self.config = config
        self.api_config = API_CONFIGS.get(api_type)
        if not self.api_config:
            raise ValueError(f"Unsupported API type: {api_type}")
        
    def get_headers(self):
        """Get headers for API request based on API type"""
        headers = {
            'User-Agent': f'InventorySlipGenerator/{APP_VERSION}',
            'Accept': 'application/json',
            'Content-Type': 'application/json'
        }
        
        api_key = self.config['API'].get(f'{self.api_type}_key')
        if not api_key:
            return headers
            
        if self.api_config['auth_type'] == 'bearer':
            headers['Authorization'] = f'Bearer {api_key}'
        elif self.api_config['auth_type'] == 'basic':
            encoded = base64.b64encode(f'{api_key}:'.encode()).decode()
            headers['Authorization'] = f'Basic {encoded}'
        
        return headers
    
    def make_request(self, endpoint, method='GET', params=None, data=None):
        """Make API request with proper error handling"""
        """Make API request with proper error handling"""
        url = f"{self.api_config['base_url']}/{self.api_config['version']}/{endpoint}"
        headers = self.get_headers()
        
        try:
            response = requests.request(
                method=method,
                url=url,
                headers=headers,
                params=params,
                json=data,
                timeout=30
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"API request failed: {str(e)}")
            raise

# Add these new routes

@app.route('/api/fetch-transfers', methods=['POST'])
@require_api_key
def fetch_transfers():
    """Fetch transfer data from selected API"""
    try:
        api_type = request.form.get('api_type', 'bamboo')
        date_from = request.form.get('date_from')
        date_to = request.form.get('date_to')
        
        config = load_config()
        client = APIClient(api_type, config)
        
        # Fetch data based on API type
        if api_type == 'bamboo':
            data = client.make_request('transfers', params={'start_date': date_from, 'end_date': date_to})
            result_df = parse_bamboo_data(data)
        elif api_type == 'cultivera':
            data = client.make_request('manifests', params={'fromDate': date_from, 'toDate': date_to})
            result_df = parse_cultivera_data(data)
        elif api_type == 'growflow':
            data = client.make_request('inventory/transfers', params={'dateStart': date_from, 'dateEnd': date_to})
            result_df = parse_growflow_data(data)
        else:
            return jsonify({'error': 'Unsupported API type'}), 400
        
        if result_df is None or result_df.empty:
            return jsonify({'error': 'No data found'}), 404
            
        # Store in session
        session['df_json'] = result_df.to_json(orient='records')
        session['format_type'] = api_type
        session['raw_json'] = json.dumps(data)
        
        return jsonify({
            'success': True,
            'message': f'Successfully fetched {len(result_df)} records',
            'redirect': url_for('data_view')
        })
        
    except Exception as e:
        logger.error(f"API fetch error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/validate-key', methods=['POST'])
def validate_api_key():
    """Validate API key for selected service"""
    try:
        api_type = request.form.get('api_type')
        api_key = request.form.get('api_key')
        
        if not api_type or not api_key:
            return jsonify({'valid': False, 'message': 'Missing required parameters'}), 400
            
        config = load_config()
        client = APIClient(api_type, {'API': {f'{api_type}_key': api_key}})
        
        # Try to make a test request
        if api_type == 'bamboo':
            client.make_request('status')
        elif api_type == 'cultivera':
            client.make_request('health')
        elif api_type == 'growflow':
            client.make_request('ping')
        
        # If we get here, the key is valid
        if 'API' not in config:
            config['API'] = {}
        config['API'][f'{api_type}_key'] = api_key
        save_config(config)
        
        return jsonify({
            'valid': True,
            'message': f'{api_type.title()} API key validated and saved'
        })
        
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 401:
            return jsonify({
                'valid': False,
                'message': 'Invalid API key'
            }), 401
        return jsonify({
            'valid': False,
            'message': f'API error: {str(e)}'
        }), e.response.status_code
    except Exception as e:
        return jsonify({
            'valid': False,
            'message': f'Validation error: {str(e)}'
        }), 500
        

@app.route('/api/settings', methods=['GET', 'POST'])
def api_settings():
    """Manage API settings"""
    config = load_config()
    
    if request.method == 'POST':
        api_type = request.form.get('api_type')
        api_key = request.form.get('api_key')
        
        if api_type and api_key:
            if 'API' not in config:
                config['API'] = {}
            config['API'][f'{api_type}_key'] = api_key
            save_config(config)
            
            flash(f'{api_type.title()} API key updated successfully')
            return redirect(url_for('settings'))
            
    # Get current API keys
    api_keys = {
        'bamboo': config.get('API', {}).get('bamboo_key', ''),
        'cultivera': config.get('API', {}).get('cultivera_key', ''),
        'growflow': config.get('API', {}).get('growflow_key', '')
    }
    
    return render_template(
        'api_settings.html',
        api_keys=api_keys,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

# Add these error handlers

class APIError(Exception):
    """Base class for API-related errors"""
    pass

class APIAuthError(APIError):
    """Authentication error"""
    pass

class APIRateLimit(APIError):
    """Rate limit exceeded"""
    pass

class APIDataError(APIError):
    """Data processing error"""
    pass

@app.errorhandler(APIError)
def handle_api_error(error):
    if isinstance(error, APIAuthError):
        flash('API authentication failed. Please check your API key.', 'error')
    elif isinstance(error, APIRateLimit):
        flash('API rate limit exceeded. Please try again later.', 'error')
    elif isinstance(error, APIDataError):
        flash('Error processing API data. Please check the format.', 'error')
    else:
        flash(f'API Error: {str(error)}', 'error')
    
    return redirect(url_for('index'))

# Add these helper functions after your imports section
def chunk_session_data(data, chunk_size=3000):
    """Split large data into smaller chunks with higher compression"""
    if not isinstance(data, str):
        data = json.dumps(data)
    
    # Use higher compression level (9 is highest)
    compressed = zlib.compress(data.encode('utf-8'), level=9)
    encoded = base64.b64encode(compressed).decode('utf-8')
    
    # Calculate optimal chunk size
    total_size = len(encoded)
    num_chunks = (total_size + chunk_size - 1) // chunk_size
    if num_chunks * 40 + total_size > 4000:  # Account for chunk metadata
        chunk_size = max(1000, (4000 - num_chunks * 40) // num_chunks)
    
    chunks = [encoded[i:i+chunk_size] for i in range(0, len(encoded), chunk_size)]
    return chunks

def store_chunked_data(key, data):
    """Store large data in session using chunks"""
    chunks = chunk_session_data(data)
    for i, chunk in enumerate(chunks):
        session[f'{key}_chunk_{i}'] = chunk
    session[f'{key}_chunks'] = len(chunks)

def get_chunked_data(key):
    """Retrieve and reconstruct chunked data from session"""
    try:
        num_chunks = session.get(f'{key}_chunks')
        if not num_chunks:
            return None
            
        chunks = []
        for i in range(num_chunks):
            chunk = session.get(f'{key}_chunk_{i}')
            if chunk is None:
                return None
            chunks.append(chunk)
        
        encoded = ''.join(chunks)
        compressed = base64.b64decode(encoded)
        decompressed = zlib.decompress(compressed)
        return json.loads(decompressed)
    except Exception as e:
        logger.error(f"Error retrieving chunked data: {str(e)}")
        return None

def clear_chunked_data(key):
    """Clear all chunks for a given key from session"""
    num_chunks = session.get(f'{key}_chunks', 0)
    for i in range(num_chunks):
        session.pop(f'{key}_chunk_{i}', None)
    session.pop(f'{key}_chunks', None)

if __name__ == '__main__':
    try:
        # Try different ports in case default is taken
        ports = [5001, 8000, 8080, 8888]
        
        for port in ports:
            try:
                # Open browser after slight delay to ensure server is running
                threading.Timer(1.5, lambda: webbrowser.open(f'http://localhost:{port}')).start()
                
                app.run(
                    host='localhost',
                    port=port,
                    debug=True,
                    use_reloader=False  # Prevent duplicate browser windows
                )
                break  # If server starts successfully, break the loop
                
            except OSError:
                if port == ports[-1]:  # If we've tried all ports
                    print("Could not find an available port. Please try again.")
                continue
                
    except Exception as e:
        print(f"Failed to start server: {str(e)}")
        sys.exit(1)
